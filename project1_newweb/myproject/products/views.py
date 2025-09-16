from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import User
from django.contrib import messages
from django.http import JsonResponse
from .form import UpdateProfileForm, UpdateUserForm, CourseForm, QuestionFileForm
from .models import Profile, Course, QuestionFile, Enrollment
from django.core.exceptions import ObjectDoesNotExist
import os
from django.conf import settings
from docx import Document
import random
import re
import logging

logger = logging.getLogger(__name__)


def user_signup(request):
    if request.method == "POST":
        username = request.POST.get('username')
        password = request.POST.get('password')
        confirm_password = request.POST.get('confirm_password')

        if not username or not password:
            messages.error(request, "Vui lòng điền đầy đủ thông tin.")
            return redirect('signup')

        if password != confirm_password:
            messages.error(request, "Mật khẩu không khớp.")
            return redirect('signup')

        if User.objects.filter(username=username).exists():
            messages.error(request, "Tên đăng nhập đã tồn tại.")
            return redirect('signup')

        user = User.objects.create_user(username=username, password=password)
       
        try:
            Profile.objects.get(user=user)
        except ObjectDoesNotExist:
            Profile.objects.create(user=user)

        messages.success(request, "Tạo tài khoản thành công. Hãy đăng nhập.")
        return redirect('login')
    return render(request, 'products/signup.html')

def user_login(request):
    if request.method == "POST":
        username = request.POST.get('username')
        password = request.POST.get('password')
        user = authenticate(request, username=username, password=password)
        if user:
            try:
                profile = user.profile  
            except Profile.DoesNotExist:
                messages.error(request, "Tên đăng nhập hoặc mật khẩu không đúng.")
                return redirect('login')

            login(request, user)
            if profile.is_admin:
                return redirect('admin_dashboard')
            messages.success(request, "Đăng nhập thành công.")
            return redirect('home')
        else:
            messages.error(request, "Tên đăng nhập hoặc mật khẩu không đúng.")
            return redirect('login')
    return render(request, 'products/login.html')

@login_required(login_url='login')
def admin_dashboard(request):
    if not request.user.profile.is_admin:
        messages.error(request, "Bạn không có quyền truy cập trang admin.")
        return redirect('home')
    
    courses = Course.objects.all()
    users = User.objects.all()
    return render(request, 'products/admin_dashboard.html', {
        'courses': courses,
        'users': users
    })

@login_required(login_url='login')
def create_admin(request):
    if request.method == "POST":
        username = request.POST.get('username')
        password = request.POST.get('password')
        confirm_password = request.POST.get('confirm_password')

        if not username or not password:
            messages.error(request, "Vui lòng điền đầy đủ thông tin.")
            return redirect('create_admin')

        if password != confirm_password:
            messages.error(request, "Mật khẩu không khớp.")
            return redirect('create_admin')

        if User.objects.filter(username=username).exists():
            messages.error(request, "Tên đăng nhập đã tồn tại.")
            return redirect('create_admin')

        user = User.objects.create_user(username=username, password=password)
        profile = Profile.objects.create(user=user, is_admin=True)
        messages.success(request, "Tạo tài khoản admin thành công.")
        return redirect('login')
    return render(request, 'products/create_admin.html')

@login_required(login_url='login')
def home(request):
    if request.user.profile.is_admin:
        return redirect('admin_dashboard')
    courses = Course.objects.all()
    return render(request, 'products/home.html', {'courses': courses})

def logout_view(request):
    logout(request)
    messages.info(request, "Bạn đã đăng xuất.")
    return redirect('login')

@login_required(login_url='login')
def profile(request):
    profile = get_object_or_404(Profile, user=request.user)
    return render(request, 'products/profile.html', {'profile': profile})

@login_required(login_url='login')
def view_user_profile(request, user_id):
    user = get_object_or_404(User, id=user_id)
    profile = Profile.objects.get(user=user)
    
    if not request.user.profile.is_admin:
        messages.error(request, "Bạn không có quyền xem hồ sơ người dùng này.")
        return redirect('admin_dashboard')
    
    owned_courses = Course.objects.filter(owner=user)
    enrolled_courses = Enrollment.objects.filter(student=user).select_related('course')
    
    return render(request, 'products/view_user_profile.html', {
        'user': user,
        'profile': profile,
        'owned_courses': owned_courses,
        'enrolled_courses': enrolled_courses,
        'is_admin': request.user.profile.is_admin
    })

@login_required(login_url='login')
def update_user(request):
    if request.method == "POST":
        uform = UpdateUserForm(request.POST, instance=request.user)
        pform = UpdateProfileForm(request.POST, request.FILES, instance=request.user.profile)
        if uform.is_valid() and pform.is_valid():
            uform.save()
            pform.save()
            messages.success(request, "Cập nhật hồ sơ thành công.")
            return redirect('profile')
        else:
            messages.error(request, "Vui lòng sửa lỗi trong biểu mẫu.")
    else:
        uform = UpdateUserForm(instance=request.user)
        pform = UpdateProfileForm(instance=request.user.profile)
    return render(request, 'products/update_user.html', {'uform': uform, 'pform': pform})

@login_required(login_url='login')
def create_course(request):
    if request.method == "POST":
        form = CourseForm(request.POST)
        if form.is_valid():
            course = form.save(commit=False)
            course.owner = request.user
            course.save()
            messages.success(request, "Tạo khóa học thành công!")
            return redirect('admin_dashboard')
        else:
            messages.error(request, "Vui lòng sửa lỗi trong biểu mẫu.")
    else:
        form = CourseForm()
    return render(request, 'products/create_course.html', {'form': form})

@login_required(login_url='login')
def course_detail(request, course_id):
    course = get_object_or_404(Course, id=course_id)
    
    
    if not request.user.profile.is_admin and request.user != course.owner and not Enrollment.objects.filter(course=course, student=request.user).exists():
       
        if not any(msg.message == "Bạn không có quyền truy cập" for msg in messages.get_messages(request)):
            logger.info(f"User {request.user.username} denied access to course {course_id}. Adding error message.")
            messages.error(request, "Bạn không có quyền truy cập")
        else:
            logger.info(f"User {request.user.username} denied access to course {course_id}. Error message already exists.")
        return redirect('home')

    question_files = QuestionFile.objects.filter(course=course).order_by('-uploaded_at')
    enrollments = Enrollment.objects.filter(course=course)

    file_content = None
    file_url = None
    file_type = None
    selected_file = None

    file_id = request.GET.get("file_id")
    if file_id:
        try:
            selected_file = QuestionFile.objects.get(id=file_id, course=course)
        except QuestionFile.DoesNotExist:
            messages.error(request, "File không tồn tại.")
    else:
        selected_file = question_files.first()

    if selected_file:
        file_path = selected_file.file.path
        file_url = selected_file.file.url

        if file_path.lower().endswith(".txt"):
            file_type = "txt"
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    file_content = f.read()
            except Exception as e:
                logger.error(f"Error reading txt file {file_path}: {str(e)}", exc_info=True)
                messages.error(request, f"Lỗi khi đọc file: {str(e)}")

        elif file_path.lower().endswith(".docx"):
            file_type = "docx"
            try:
                doc = Document(file_path)
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                file_content = "\n".join(paragraphs)
            except Exception as e:
                logger.error(f"Error reading docx file {file_path}: {str(e)}", exc_info=True)
                messages.error(request, f"Lỗi khi đọc file Word: {str(e)}")

        elif file_path.lower().endswith(".pdf"):
            file_type = "pdf"
            messages.info(request, "File PDF không hiển thị nội dung trực tiếp. Vui lòng tải xuống.")

        else:
            messages.warning(request, "Định dạng file không được hỗ trợ.")

    return render(request, "products/course_detail.html", {
        "course": course,
        "question_files": question_files,
        "enrollments": enrollments,
        "file_content": file_content,
        "file_url": file_url,
        "file_type": file_type,
        "selected_file": selected_file,
        "is_admin": request.user.profile.is_admin
    })
def _split_sentences(text):
    if not text or len(text) < 20:  
        return []
    import re
    sentences = re.split(r'(?<=[.!?])\s+', text.strip())
    return [s.strip() for s in sentences if s.strip() and len(s.strip()) > 20]

def _pick_keyword(sentence):
    words = re.findall(r"[A-Za-zÀ-ỹ0-9_']+", sentence)
    words = [w for w in words if len(w)>4 and not w.isdigit()]
    if not words:
        words = re.findall(r"[A-Za-zÀ-ỹ0-9_']+", sentence)
    if not words:
        return None
    kw = max(words, key=len)
    return kw

def generate_quiz_from_text(file_path, num_questions=None):
    questions = []
    doc = Document(file_path)
    
   
    full_text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
    if not full_text:
        return questions
    
   
    question_blocks = re.split(r'\n(?=\d+\.\s)', full_text.strip())
    
    for block in question_blocks:
        if not block.strip():
            continue
        
       
        question_match = re.match(r'^(\d+\.\s*.*?)\n((?:[A-D]\..*?\n)*)(?:Đáp án:\s*([A-D])|$)', block, re.DOTALL)
        if not question_match:
            logger.warning(f"Không khớp regex với block: {block}")
            continue
        
        question_text = question_match.group(1).strip()
        options_text = question_match.group(2).strip() if question_match.group(2) else ""
        correct_answer = question_match.group(3) 
        
        
        options = re.findall(r'([A-D]\.[^\n]*?)(?=\n[A-D]\.|$)', options_text, re.DOTALL)
        options = [opt.strip() for opt in options]
        if len(options) < 4:
            logger.warning(f"Chỉ tìm thấy {len(options)} options trong block: {options_text}")
            continue
           
        correct_indices = []
        seen_texts = set()  
        content_to_index = {}
        for idx, opt in enumerate(options):
            content = opt.split('.', 1)[1].strip() if '.' in opt else opt.strip()
            content_to_index[content] = idx
            content_to_index[opt.strip()] = idx
        
        for p in doc.paragraphs:
            for run in p.runs:
                if run.bold and run.text.strip() and run.text.strip() not in seen_texts:
                    bold_text = run.text.strip()
                    seen_texts.add(bold_text)
                    
                    
                    found = False
                    for content, idx in content_to_index.items():
                        if bold_text == content or bold_text in content:
                            if idx not in correct_indices:
                                correct_indices.append(idx)
                                logger.info(f"Found bold answer at index {idx}: {bold_text}")
                                found = True
                                break
                    
                    if not found:
                        for content, idx in content_to_index.items():
                            if content in bold_text:
                                if idx not in correct_indices:
                                    correct_indices.append(idx)
                                    logger.info(f"Found reverse bold match at index {idx}: {bold_text} contains {content}")
                                    break
        

        if correct_answer and not correct_indices:
            correct_index = ord(correct_answer.upper()) - ord('A')
            if correct_index < len(options):
                correct_indices.append(correct_index)
                logger.info(f"Using Đáp án: {correct_answer} at index {correct_index}")
        

        if correct_indices:
            questions.append({
                'question': question_text,
                'options': options[:4], 
                'answer_indices': correct_indices,
                'hint': f"Đáp án: {correct_answer if correct_answer else 'Được đánh dấu in đậm'}"
            })
        else:
            logger.warning(f"Không tìm thấy đáp án cho câu hỏi: {question_text[:50]}...")
        
        if len(questions) >= num_questions:
            break
    
    logger.info(f"Đã tạo {len(questions)} câu hỏi")
    return questions[:num_questions]

@login_required(login_url='login')
def quiz_from_file(request, course_id, file_id):
    course = get_object_or_404(Course, id=course_id)
    try:
        qfile = QuestionFile.objects.get(id=file_id, course=course)
    except QuestionFile.DoesNotExist:
        messages.error(request, "File không tồn tại.")
        return redirect('course_detail', course_id=course.id)

    file_path = qfile.file.path
    logger.info(f"Processing file: {file_path}")  
    
    questions = generate_quiz_from_text(file_path, num_questions=10) 
    if not questions:
        logger.warning("Không tạo được questions từ file content.")
        messages.warning(request, "Không thể tạo quiz từ file. Vui lòng kiểm tra định dạng file.")
        return redirect('course_detail', course_id=course.id)

    enumerated_questions = [
        (i, {
            'question': q['question'],
            'options': list(enumerate(q['options'], start=0)),
            'answer_indices': q['answer_indices'],
            'hint': q['hint']
        })
        for i, q in enumerate(questions)
    ]

    file_basename = os.path.basename(qfile.file.name)
    logger.info(f"Tạo quiz thành công với {len(questions)} câu hỏi cho file {file_basename}")

    if request.method == 'POST':
        score = 0
        total_questions = len(questions)
        results = []
        logger.info(f"POST data: {request.POST.dict()}")
        for i, q in enumerate(questions):
            selected = request.POST.getlist(f'q{i}')
            selected = [int(s) for s in selected if s.isdigit()]  
            correct_indices = q['answer_indices']
            logger.info(f"Question {i}: Selected {selected}, Correct {correct_indices}, Options {q['options']}")
            is_correct = False
            if correct_indices:  
                is_correct = any(idx in correct_indices for idx in selected)
            if is_correct:
                score += 1
            enumerated_options = list(enumerate(q['options'], start=0))
            results.append({
                'question': q['question'],
                'options': enumerated_options,
                'selected': selected,
                'is_correct': is_correct,
                'correct_indices': correct_indices
            })
        return render(request, 'products/quiz_from_file.html', {
            'course': course,
            'file': qfile,
            'enumerated_questions': enumerated_questions,
            'is_admin': request.user.profile.is_admin,
            'results': results,
            'score': score,
            'total_questions': total_questions,
            'questions': questions  
        })

    return render(request, 'products/quiz_from_file.html', {
        'course': course,
        'file': qfile,
        'enumerated_questions': enumerated_questions,
        'is_admin': request.user.profile.is_admin,
        'questions': questions  
    })
@login_required(login_url='login')
def flashcards_from_file(request, course_id, file_id):
    course = get_object_or_404(Course, id=course_id)
    try:
        qfile = QuestionFile.objects.get(id=file_id, course=course)
    except QuestionFile.DoesNotExist:
        messages.error(request, "File không tồn tại.")
        return redirect('course_detail', course_id=course.id)

    file_path = qfile.file.path
    cards = []
    
    if file_path.lower().endswith('.docx'):
        try:
            doc = Document(file_path)
            logger.info(f"Processing flashcards from file: {file_path}")
            
            
            full_text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
            logger.info(f"File content:\n{full_text}\n")
            
            if not full_text:
                messages.error(request, "File Word trống hoặc không chứa nội dung hợp lệ.")
                return redirect('course_detail', course_id=course.id)
            
          
            question_blocks = re.split(r'\n(?=\d+\.\s)', full_text.strip(), flags=re.MULTILINE)
            
            for block in question_blocks:
                if not block.strip():
                    continue
                
               
                question_match = re.match(
                    r'^(\d+\.\s*.*?)(?=\n[A-D]\.|$)\n((?:[A-D]\..*?\n)*)?(?:Đáp án:\s*([A-D])\s*$|$)',
                    block, re.DOTALL | re.IGNORECASE
                )
                if not question_match:
                    logger.warning(f"Regex did not match for block:\n{block[:50]}...")
                    continue
                
                question_text = question_match.group(1).strip()
                options_text = question_match.group(2).strip() if question_match.group(2) else ""
                correct_answer_key = question_match.group(3)
                
                # Trích xuất các lựa chọn (A., B., C., D.)
                options = re.findall(r'([A-D]\.[^\n]*?)(?=\n[A-D]\.|$|\nĐáp án:)', options_text, re.DOTALL)
                options = [opt.strip() for opt in options]
                if len(options) < 4:
                    logger.warning(f"Only found {len(options)} options in block:\n{options_text[:50]}...")
                    continue
                
               
                correct_answer = None
                if correct_answer_key:
                    correct_index = ord(correct_answer_key.upper()) - ord('A')
                    if correct_index < len(options):
                        correct_answer = options[correct_index]
                        logger.info(f"Found flashcard - Q: {question_text[:50]}... A: {correct_answer}")
                
                if correct_answer:
                    cards.append({
                        'front': question_text,
                        'back': correct_answer
                    })
                else:
                    logger.warning(f"No valid answer found for question: {question_text[:50]}...")
            
            logger.info(f"Created {len(cards)} flashcards from Word file")
            
        except Exception as e:
            logger.error(f"Error processing Word file: {str(e)}", exc_info=True)
            messages.error(request, f"Lỗi khi đọc file Word: {str(e)}")
            return redirect('course_detail', course_id=course.id)
    
    elif file_path.lower().endswith('.txt'):
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                file_content = f.read()
            
            logger.info(f"Processing flashcards from text file: {file_path}")
            logger.info(f"File content:\n{file_content}\n")
            
            
            question_blocks = re.split(r'\n(?=\d+\.\s)', file_content.strip(), flags=re.MULTILINE)
            
            for block in question_blocks:
                if not block.strip():
                    continue
                
                question_match = re.match(
                    r'^(\d+\.\s*.*?)(?=\n[A-D]\.|$)\n((?:[A-D]\..*?\n)*)?(?:Đáp án:\s*([A-D])\s*$|$)',
                    block, re.DOTALL | re.IGNORECASE
                )
                if not question_match:
                    logger.warning(f"Regex did not match for block:\n{block[:50]}...")
                    continue
                
                question_text = question_match.group(1).strip()
                options_text = question_match.group(2).strip() if question_match.group(2) else ""
                correct_answer_key = question_match.group(3)
                
                options = re.findall(r'([A-D]\.[^\n]*?)(?=\n[A-D]\.|$|\nĐáp án:)', options_text, re.DOTALL)
                options = [opt.strip() for opt in options]
                if len(options) < 4:
                    logger.warning(f"Only found {len(options)} options in block:\n{options_text[:50]}...")
                    continue
                
                correct_answer = None
                if correct_answer_key:
                    correct_index = ord(correct_answer_key.upper()) - ord('A')
                    if correct_index < len(options):
                        correct_answer = options[correct_index]
                        logger.info(f"Found flashcard - Q: {question_text[:50]}... A: {correct_answer}")
                
                if correct_answer:
                    cards.append({
                        'front': question_text,
                        'back': correct_answer
                    })
                else:
                    logger.warning(f"No valid answer found for question: {question_text[:50]}...")
            
            logger.info(f"Created {len(cards)} flashcards from text file")
                
        except Exception as e:
            logger.error(f"Error processing text file: {str(e)}", exc_info=True)
            messages.error(request, f"Lỗi khi đọc file: {str(e)}")
            return redirect('course_detail', course_id=course.id)
    
    else:
        messages.warning(request, "Flashcards hiện chỉ hỗ trợ file .txt và .docx. Với PDF vui lòng tải về và chuyển đổi sang .docx hoặc .txt.")
        return redirect('course_detail', course_id=course.id)

    if not cards:
        messages.error(request, "Không tìm thấy nội dung phù hợp để tạo flashcards. Hãy đảm bảo file có câu hỏi với số thứ tự (VD: 1.), 4 lựa chọn (A., B., C., D.), và dòng 'Đáp án: [A-D]'.")
        return redirect('course_detail', course_id=course.id)

    return render(request, 'products/flashcards_from_file.html', {
        'course': course,
        'file': qfile,
        'cards': cards,
        'is_admin': request.user.profile.is_admin
    })


@login_required(login_url='login')
def delete_question_file(request, course_id, file_id):
    course = get_object_or_404(Course, id=course_id)
    file = get_object_or_404(QuestionFile, id=file_id, course=course)
    
    if request.user != course.owner and not request.user.profile.is_admin:
        messages.error(request, "Bạn không có quyền xóa tài liệu này.")
        return redirect('course_detail', course_id=course.id)

    if file.file:
        file.file.delete(save=False)
    file.delete()
    messages.success(request, "Xóa tài liệu thành công!")
    return redirect("course_detail", course_id=course.id)

@login_required(login_url='login')
def upload_question_file(request, course_id):
    course = get_object_or_404(Course, id=course_id)
    
    if request.user != course.owner and not request.user.profile.is_admin:
        logger.warning(f"User {request.user.username} attempted to upload file to course {course.id} without permission.")
        return JsonResponse({'success': False, 'message': "Bạn không có quyền upload file."}, status=403)

    if request.method == "POST":
        form = QuestionFileForm(request.POST, request.FILES)
        if form.is_valid():
            file = request.FILES.get('file')
            if not file:
                logger.error("No file provided in the request.")
                return JsonResponse({'success': False, 'message': "Không tìm thấy file."}, status=400)

            if not file.name.lower().endswith('.docx'):
                logger.error(f"Invalid file format: {file.name}")
                return JsonResponse({'success': False, 'message': "Chỉ được upload file Word (.docx)."}, status=400)

            original_name = file.name
            counter = 1
            base_name, ext = os.path.splitext(original_name)
            file_name = original_name
            while QuestionFile.objects.filter(course=course, file=f'question_files/{file_name}').exists():
                file_name = f"{base_name}_{counter}{ext}"
                counter += 1
            file.name = file_name

            try:
                qfile = form.save(commit=False)
                qfile.course = course
                qfile.uploaded_by = request.user
                qfile.save()
                logger.info(f"File {file.name} uploaded successfully for course {course.id} by user {request.user.username}.")
                return JsonResponse({'message': "đã upfile thành công"}, status=200)
            except Exception as e:
                logger.error(f"Error saving file {file.name}: {str(e)}", exc_info=True)
                return JsonResponse({'success': False, 'message': f"Lỗi khi lưu file: {str(e)}"}, status=500)
        else:
            logger.error(f"Form errors: {form.errors}")
            return JsonResponse({'success': False, 'errors': form.errors.as_json()}, status=400)
    else:
        form = QuestionFileForm()

    return render(request, 'products/upload_question_file.html', {'form': form, 'course': course})

@login_required(login_url='login')
def manage_enrollments(request, course_id):
    course = get_object_or_404(Course, id=course_id)
    if request.user != course.owner and not request.user.profile.is_admin:
        messages.error(request, "Bạn không có quyền quản lý học viên của khóa học này.")
        return redirect('home')

    if request.method == "POST":
        student_username = request.POST.get('student_username')
        try:
            student = User.objects.get(username=student_username)
            if Enrollment.objects.filter(course=course, student=student).exists():
                messages.error(request, "Học viên này đã được ghi danh.")
            else:
                Enrollment.objects.create(course=course, student=student)
                messages.success(request, "Thêm học viên thành công!")
        except User.DoesNotExist:
            messages.error(request, "Không tìm thấy học viên với tên này.")
        return redirect('manage_enrollments', course_id=course.id)

    enrollments = course.enrollments.all()
    return render(request, 'products/manage_enrollments.html', {
        'course': course,
        'enrollments': enrollments,
        'is_admin': request.user.profile.is_admin
    })


@login_required(login_url='login')
def delete_course(request, course_id):
    if request.method != 'POST':
        return JsonResponse({'success': False, 'message': 'Phương thức không được hỗ trợ.'}, status=405)
    if not request.user.profile.is_admin:
        return JsonResponse({'success': False, 'message': 'Bạn không có quyền xóa khóa học.'}, status=403)
    
    course = get_object_or_404(Course, id=course_id)
    try:
        for qfile in QuestionFile.objects.filter(course=course):
            if qfile.file and os.path.isfile(qfile.file.path):
                try:
                    os.remove(qfile.file.path)
                    logger.info(f"Deleted file: {qfile.file.path}")
                except OSError as e:
                    logger.error(f"Failed to delete file {qfile.file.path}: {str(e)}")
                    return JsonResponse({'success': False, 'message': f'Không thể xóa file {qfile.file.name}: {str(e)}'}, status=500)
            else:
                logger.warning(f"File not found: {qfile.file.path}")
        course.delete()
        logger.info(f"Successfully deleted course {course_id}")
        return JsonResponse({'success': True, 'message': 'Xóa khóa học thành công!'}, status=200)
    except Exception as e:
        logger.error(f"Error deleting course {course_id}: {str(e)}", exc_info=True)
        return JsonResponse({'success': False, 'message': f'Lỗi khi xóa khóa học: {str(e)}'}, status=500)